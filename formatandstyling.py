# =======================
# formatting_pipeline.py
# =======================

# Imports & Environment
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from langchain_groq import ChatGroq
from dotenv import load_dotenv
import os
import json
from pathlib import Path

# Correct way to create output file path


# Load environment variables
load_dotenv()
GROQ_API_KEY = os.getenv("GROQ_API_KEY")

# Initialize LLM
llm = ChatGroq(
    model="llama3-8b-8192",
    temperature=0,
    groq_api_key=GROQ_API_KEY
)
import json
import re

def clean_llm_json_output(llm_output: str) -> str:
    """
    Cleans LLM output to extract valid JSON string.
    """
    # Remove Markdown code fences if present
    cleaned = re.sub(r"```(json)?", "", llm_output, flags=re.IGNORECASE)
    # Remove any leading/trailing text before/after JSON
    match = re.search(r"\{.*\}", cleaned, flags=re.DOTALL)
    if match:
        return match.group()
    else:
        # fallback: empty JSON structure
        return '{"sections":[]}'

# =======================
# Template Configurations
# =======================
TEMPLATES = {
    "report": {
        "font": "Times New Roman",
        "font_size": 12,
        "heading1_size": 16,
        "heading2_size": 14,
        "toc": True,
        "logo": "company_logo.png"
    },
    "proposal": {
        "font": "Arial",
        "font_size": 12,
        "heading1_size": 16,
        "heading2_size": 14,
        "toc": True,
        "logo": "company_logo.png"
    },
    "resume": {
        "font": "Calibri",
        "font_size": 11,
        "heading1_size": 14,
        "toc": False
    },
    "meeting_notes": {
        "font": "Calibri",
        "font_size": 12,
        "heading1_size": 14,
        "toc": False
    },
        "other": {
        "font": "Arial",
        "font_size": 12,
        "heading1_size": 16,
        "heading2_size": 14,
        "toc": False
    }
}

# =======================
# LLM Functions
# =======================
def detect_document_type(doc_text: str) -> str:
    """
    Detect document type using LLM.
    """
    prompt = f"""
    You are an expert document classifier.
    Classify the following document into one of: 'report', 'proposal', 'resume', 'meeting_notes', 'other'.
    Document Text:
    \"\"\"{doc_text}\"\"\"
    """
    response = llm.invoke(prompt)
    return response.content.strip() if hasattr(response, "content") else response.strip()

def generate_document_outline(doc_text: str) -> str:
    """
    Generate structured outline from text using LLM.
    Output is a JSON string with sections, subheadings, and paragraphs or sentences and if there are any key points or takeways write them in bullet points.
    """
    prompt = f"""
    You are an expert document designer.
    Analyze the following text and generate a structured outline in JSON format:
    - sections: list of sections
    - each section has 'heading', 'subheadings', each with 'heading' paragraph ,sentence and 'bullets'

    Document Text:
    \"\"\"{doc_text}\"\"\"
    Only return **valid JSON**, no extra text.
    """
    response = llm.invoke(prompt)
    print(response.content)
    return response.content if hasattr(response, "content") else response

# =======================
# Document Loader
# =======================
def load_document(file_path: str):
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".docx":
        from langchain_community.document_loaders import Docx2txtLoader
        loader = Docx2txtLoader(file_path)
    elif ext == ".pdf":
        from langchain_community.document_loaders import PyPDFLoader
        loader = PyPDFLoader(file_path)
    else:
        raise ValueError("Unsupported format for this demo")
    return loader.load()

# =======================
# Rendering Functions
# =======================
def render_outline_to_docx(outline_json: str, template_config: dict, output_file: str):
    doc = Document()

    # Add logo if exists
    if "logo" in template_config:
        try:
            doc.add_picture(template_config["logo"], width=Inches(1.5))
        except:
            pass

    # Clean and parse JSON
    cleaned_json = clean_llm_json_output(outline_json)
    try:
        outline_data = json.loads(cleaned_json)
    except json.JSONDecodeError:
        outline_data = {"sections": [{"heading": "Document", "subheadings": [], "bullets": []}]}

    # Add sections and bullets
    for section in outline_data.get("sections", []):
        doc.add_heading(section.get("heading", ""), level=1)
        for sub in section.get("subheadings", []):
            doc.add_heading(sub.get("heading", ""), level=2)
            for bullet in sub.get("bullets", []):
                doc.add_paragraph(f"- {bullet}", style="List Bullet")

    # Apply fonts and styles
    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith("Heading 1"):
            paragraph.style.font.name = template_config.get("font", "Arial")
            paragraph.style.font.size = Pt(template_config.get("heading1_size", 16))
        elif paragraph.style.name.startswith("Heading 2"):
            paragraph.style.font.name = template_config.get("font", "Arial")
            paragraph.style.font.size = Pt(template_config.get("heading2_size", 14))
        else:
            paragraph.style.font.name = template_config.get("font", "Arial")
            paragraph.style.font.size = Pt(template_config.get("font_size", 12))
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY



    doc.save(output_file)
    return output_file

# =======================
# Main Pipeline
# =======================
def formatting_pipeline(file_path: str) -> str:
    # Load document and extract text
    docs = load_document(file_path)
    full_text = " ".join([doc.page_content for doc in docs])

    # Detect document type
    doc_type = detect_document_type(full_text)
    template_config = TEMPLATES.get(doc_type, TEMPLATES["report"])

    # Generate outline
    outline = generate_document_outline(full_text)

    # Render formatted Word doc
    output_file = file_path.with_name(file_path.stem + "_formatted.docx")
    render_outline_to_docx(outline, template_config, output_file)
    return output_file

# =======================
# Example Usage
# =======================
if __name__ == "__main__":
    input_file = "project.docx"
    formatted_file = formatting_pipeline(input_file)
    print(f"Formatted document saved at: {formatted_file}")
